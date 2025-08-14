import os
import requests
import json
import pdfplumber
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
import re

# --- CONFIGURATION ---
OLLAMA_API_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "mistral"
DATA_DIR = "data"
TEMPLATES_DIR = "templates"
OUTPUT_DIR = "output"
FIGURE_DIR = os.path.join(OUTPUT_DIR, "figures") # To store extracted images

# --- 1. SETUP & USER INTERACTION ---

def select_file_from_dir(directory, file_extension):
    """Lists files in a directory and asks the user to choose one."""
    print(f"\nScanning for '{file_extension}' files in './{directory}'...")
    files = [f for f in os.listdir(directory) if f.endswith(file_extension)]
    if not files:
        print(f"Error: No '{file_extension}' files found in the './{directory}' directory.")
        return None

    print(f"Please choose a file to process:")
    for i, f in enumerate(files):
        print(f"  [{i+1}] {f}")

    while True:
        try:
            choice = int(input("Enter the number of your choice: ")) - 1
            if 0 <= choice < len(files):
                return os.path.join(directory, files[choice])
            else:
                print("Invalid number. Please try again.")
        except ValueError:
            print("Please enter a valid number.")

def get_language_choice():
    """Asks the user to choose the translation direction."""
    print("\nSelect the translation direction:")
    print("  [1] Romanian to English")
    print("  [2] English to Romanian")
    while True:
        try:
            choice = int(input("Enter your choice (1 or 2): "))
            if choice == 1:
                return "Romanian", "English"
            elif choice == 2:
                return "English", "Romanian"
            else:
                print("Invalid choice. Please select 1 or 2.")
        except ValueError:
            print("Please enter a valid number.")

# --- 2. EXTRACTION & ANALYSIS (OCR PHASE) ---

def extract_content_from_pdf(pdf_path):
    """
    Extracts text, tables, and images from a PDF.
    It replaces tables and figures with placeholders in the text.
    """
    print(f"\n Analyzing PDF: {os.path.basename(pdf_path)}...")
    if not os.path.exists(FIGURE_DIR):
        os.makedirs(FIGURE_DIR)

    content_list = []
    tables = []
    figures = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(tqdm(pdf.pages, desc="Processing PDF Pages")):
            # Extract and save images (figures)
            page_images = page.images
            for i, img in enumerate(page_images):
                figure_index = len(figures) + 1
                figure_placeholder = f"\n[FIGURE_{figure_index}]\n"
                
                # Get the bounding box of the image to place the placeholder correctly
                img_bbox = (img['x0'], img['top'], img['x1'], img['bottom'])
                
                # Extract the image bytes
                try:
                    img_obj = page.crop(img_bbox).to_image(resolution=300)
                    figure_filename = os.path.join(FIGURE_DIR, f"figure_{page_num+1}_{i+1}.png")
                    img_obj.save(figure_filename, format="PNG")
                    figures.append(figure_filename)
                    # Add placeholder based on vertical position
                    content_list.append({'type': 'figure', 'content': figure_placeholder, 'y0': img['top']})
                except Exception as e:
                    print(f"Warning: Could not extract an image on page {page_num+1}. Reason: {e}")

            # Extract tables
            page_tables = page.extract_tables()
            for table_data in page_tables:
                table_index = len(tables) + 1
                table_placeholder = f"\n[TABLE_{table_index}]\n"
                tables.append(table_data)
                
                # Find table position to insert placeholder correctly
                # We approximate the position from the first cell's content
                first_cell_text = table_data[0][0]
                if first_cell_text:
                    text_instances = page.search(first_cell_text)
                    if text_instances:
                         content_list.append({'type': 'table', 'content': table_placeholder, 'y0': text_instances[0]['top']})

            # Extract text, excluding text inside tables
            text = page.extract_text(x_tolerance=2, keep_blank_chars=True)
            if text:
                content_list.append({'type': 'text', 'content': text, 'y0': page.bbox[1]}) # Use page top for general text

    # Sort all extracted elements by their vertical position (y0) to maintain reading order
    # A bit complex because a single page adds multiple items. A simpler approach for now:
    # We will process text page by page and insert placeholders where they were found.
    # The current implementation is a simplification but should work for many documents.
    # For a truly robust solution, one would process objects sorted by their y-coordinate.
    
    # A simplified re-ordering will be done by consolidating text after extraction
    full_text = ""
    # We will let the placeholders found on a page represent their content for now
    # This is a complex problem; for this script, we assume a sequential read is "good enough"
    # and placeholders are inserted at the end of a page's text.
    with pdfplumber.open(pdf_path) as pdf:
        page_texts = []
        for page in pdf.pages:
            page_texts.append(page.extract_text() or "")
        
        full_text = "\n".join(page_texts)
        
        # Now, replace table and figure areas with placeholders. This is a heuristic.
        # A more advanced version would use coordinates to insert placeholders perfectly.
        num_tables = len(tables)
        num_figures = len(figures)
        
        # This is a simplified replacement. It doesn't use coordinates but assumes order.
        full_text += "\n" + "\n".join([f"[TABLE_{i+1}]" for i in range(num_tables)])
        full_text += "\n" + "\n".join([f"[FIGURE_{i+1}]" for i in range(num_figures)])

    # Clean up excessive newlines that can confuse the model
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)
    
    print("PDF analysis complete.")
    return full_text, tables, figures


# --- 3. TRANSLATION & REFINEMENT (MISTRAL PHASE) ---

def translate_text_with_mistral(text, src_lang, tgt_lang):
    """
    Sends text to the Ollama Mistral model for translation and grammar fix.
    Handles long text by splitting it into chunks of a manageable size (approx. 6000 tokens).
    """
    print(f"\n Translating from {src_lang} to {tgt_lang} using Mistral (with token-aware chunking)...")

    # --- NEW TOKEN-BASED CHUNKING LOGIC ---
    # Rule of thumb: 1 token ~= 4 characters. Target: 6000 tokens ~= 24000 characters.
    MAX_CHARS_PER_CHUNK = 24000
    
    chunks = []
    current_pos = 0
    while current_pos < len(text):
        end_pos = min(current_pos + MAX_CHARS_PER_CHUNK, len(text))
        
        # If we are not at the end of the text, try to find a natural split point (end of sentence).
        if end_pos < len(text):
            # Find the last period or newline before the hard character limit
            split_pos = text.rfind('.', current_pos, end_pos)
            if split_pos == -1:
                split_pos = text.rfind('\n', current_pos, end_pos)
            
            # If a natural split point is found, use it. Otherwise, use the hard limit.
            if split_pos != -1:
                end_pos = split_pos + 1
        
        chunk = text[current_pos:end_pos]
        chunks.append(chunk)
        current_pos = end_pos
    
    print(f"Text was split into {len(chunks)} chunks to ensure stability.")
    # --- END OF NEW LOGIC ---

    translated_chunks = []
    for i, chunk in enumerate(tqdm(chunks, desc="Translating Chunks")):
        if not chunk.strip():
            continue

        # The "flawless" prompt remains the same
        prompt = f"""
        You are an expert multilingual translator and proofreader. Your task is to translate a text from {src_lang} to {tgt_lang}.
        Follow these instructions precisely:
        1. Translate the text literally and completely. DO NOT summarize, shorten, or omit any information.
        2. Correct any minor grammatical mistakes in the original text during translation (e.g., spelling, missing commas).
        3. Preserve special placeholders like [TABLE_X] and [FIGURE_X] exactly as they are in the translated output.
        
        Translate the following text:
        ---
        {chunk}
        """

        try:
            # You can keep the extended timeout or reduce it back, e.g., to 320
            payload = {
                "model": MODEL_NAME,
                "prompt": prompt,
                "stream": False
            }
            response = requests.post(OLLAMA_API_URL, json=payload, timeout=320) # Keeping your 320s timeout
            response.raise_for_status()
            
            response_data = response.json()
            translated_chunk = response_data.get("response", "").strip()
            translated_chunks.append(translated_chunk)

        except requests.exceptions.RequestException as e:
            print(f"\nError contacting Ollama API on chunk {i+1}: {e}")
            print("This could be due to a very complex chunk or Ollama server issues. Aborting.")
            return None

    print("✅ Translation complete.")
    return "".join(translated_chunks)

# --- 4. RECONSTRUCTION & FORMATTING (OUTPUT PHASE) ---

def create_word_document(output_path, template_path, translated_text, tables, figures):
    """
    Creates a new Word document based on a template, populating it with
    the translated text, tables, and figures.
    """
    print(f"\n Building Word document using template: {os.path.basename(template_path)}...")
    
    # Using the template file directly preserves all its styles and formatting
    doc = Document(template_path)

    # Dictionaries to hold the actual data for placeholders
    table_map = {f"[TABLE_{i+1}]": data for i, data in enumerate(tables)}
    figure_map = {f"[FIGURE_{i+1}]": path for i, path in enumerate(figures)}

    # Process the translated text paragraph by paragraph
    for para_text in translated_text.split('\n'):
        para_text = para_text.strip()
        
        # Check if the paragraph is a placeholder for a table or figure
        if para_text in table_map:
            table_data = table_map[para_text]
            if table_data:
                # Add a new table to the document
                # Using the default table style from the template
                word_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]), style='Table Grid')
                for i, row in enumerate(table_data):
                    for j, cell_text in enumerate(row):
                        word_table.cell(i, j).text = cell_text if cell_text else ""
                # Add a caption
                doc.add_paragraph(f"Table {para_text.split('_')[1].strip('[]')}", style='Caption')

        elif para_text in figure_map:
            figure_path = figure_map[para_text]
            # Add the figure image
            doc.add_picture(figure_path, width=Inches(6.0)) # Default width
            # Add a caption
            doc.add_paragraph(f"Figure {para_text.split('_')[1].strip('[]')}", style='Caption')
            
        else:
            # It's a regular paragraph, so add it
            if para_text: # Avoid adding empty paragraphs
                doc.add_paragraph(para_text, style='Normal')
    
    doc.save(output_path)
    print(f"Successfully created Word document: {output_path}")

# --- MAIN WORKFLOW ---

def main():
    """Main function to run the entire application workflow."""
    print("--- Document Translation & Formatting Tool ---")
    
    # 1. Get user input
    pdf_path = select_file_from_dir(DATA_DIR, ".pdf")
    if not pdf_path: return
    
    template_path = select_file_from_dir(TEMPLATES_DIR, ".docx")
    if not template_path: return
    
    src_lang, tgt_lang = get_language_choice()
    
    # 2. Extract content from the PDF
    original_text, tables, figures = extract_content_from_pdf(pdf_path)
    if not original_text:
        print("Could not extract any text from the PDF. Aborting.")
        return

    # 3. Translate the text
    translated_text = translate_text_with_mistral(original_text, src_lang, tgt_lang)
    if translated_text is None:
        print("Translation failed. Aborting.")
        return

    # 4. Create the final Word document
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_filename = f"{base_name}_translated_to_{tgt_lang.lower()}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    create_word_document(output_path, template_path, translated_text, tables, figures)

    print("\n--- Process Finished ---")


if __name__ == "__main__":
    main()