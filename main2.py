import os
import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Define paths
base_path = r"C:\Users\530985\Desktop\App_developing\DocumentHandler"
data_folder = os.path.join(base_path, "data")
output_folder = os.path.join(base_path, "output")

# Ensure output folder exists
os.makedirs(output_folder, exist_ok=True)

# Detect and crop diagram-like regions
def crop_diagram_regions(pil_image):
    cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edged = cv2.Canny(blurred, 50, 150)

    contours, _ = cv2.findContours(edged, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    cropped_images = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 100 and h > 100:  # Filter small regions
            cropped = pil_image.crop((x, y, x + w, y + h))
            cropped_images.append(cropped)

    return cropped_images

# Extract text and diagram crops from PDF
def extract_text_and_diagram_crops(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = ""
    diagram_crops = []

    for page in doc:
        full_text += page.get_text() + "\n"

        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(BytesIO(image_bytes))

            crops = crop_diagram_regions(image)
            diagram_crops.extend(crops)

    return full_text, diagram_crops

# Save text and cropped diagrams to DOCX
def save_to_docx(text, diagrams, output_path):
    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    doc.add_paragraph(text)

    if diagrams:
        doc.add_heading("Extracted Diagrams", level=1)
        for idx, img in enumerate(diagrams):
            temp_img_path = os.path.join(output_folder, f"temp_crop_{idx}.png")
            img.save(temp_img_path)
            doc.add_picture(temp_img_path, width=Inches(5))
            os.remove(temp_img_path)

    doc.save(output_path)

# Process PDFs
for filename in os.listdir(data_folder):
    if filename.lower().endswith(".pdf"):
        pdf_path = os.path.join(data_folder, filename)
        text, diagram_crops = extract_text_and_diagram_crops(pdf_path)

        output_filename = os.path.splitext(filename)[0] + "_cropped_diagrams.docx"
        output_path = os.path.join(output_folder, output_filename)
        save_to_docx(text, diagram_crops, output_path)

print("Cropped diagrams and text saved to DOCX in the output folder.")
