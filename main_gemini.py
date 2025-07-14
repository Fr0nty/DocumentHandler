import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from langchain import Gemini  # Import your Google Gemini library
from docx import Document
from docx.shared import Inches
import io

# Function to read the API key from a file
def read_api_key(filepath):
    with open(filepath, 'r') as file:
        return file.read().strip()

# Function to authenticate and use Gemini API
def gemini_api_auth(api_key):
    return Gemini(api_key)

# Step 1: Extract Text, Figures, and Tables
def extract_text_from_pdf(pdf_path):
    print("Starting text extraction from PDF...")
    doc = fitz.open(pdf_path)
    text = ""
    figures = {}
    tables = {}
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
        print(f"Extracted text from page {page_num + 1}")

        # Detect and extract images
        images = page.get_image_list(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            figures[f"Figure_{page_num + 1}_{img_index + 1}"] = image
            print(f"Extracted figure {img_index + 1} from page {page_num + 1}")

        # Add table extraction analogously if needed 

    print("Completed text extraction from PDF")
    return text, figures, tables

# Step 2: Cleaning and Translation using Gemini
def correct_text(text, gemini):
    print("Starting grammar correction...")
    corrected_text = gemini.correct(text)
    print("Completed grammar correction")
    return corrected_text

def translate_text(text, gemini, target_language="en"):
    print(f"Starting translation to {target_language}...")
    translated_text = gemini.translate(text, target_language=target_language)
    print("Completed translation")
    return translated_text

# Step 3: Reinserting Labels
def label_figures_text(text, figures):
    print("Starting reinsertion of figure labels...")
    figure_labels = "\n"  # Collect all figure labels to add to the end of the document
    for label, image in figures.items():
        figure_labels += f"{label}\n"
        print(f"Labeled {label}")
    print("Completed reinsertion of figure labels")
    return text + figure_labels

# Step 4: Creating Word Document According to Template
def create_word_document(template_path, text, figures, output_path):
    print("Starting Word document creation...")
    document = Document(template_path)

    # Extract template styles (if needed)
    # Add text
    paragraphs = text.split('\n')
    for para in paragraphs:
        document.add_paragraph(para)
    
    # Add Figures
    for label, image in figures.items():
        document.add_heading(label, level=2)
        image_path = f"{label}.png"
        image.save(image_path)
        document.add_picture(image_path, width=Inches(5.0))  # Adjust size as needed

    document.save(output_path)
    print("Completed Word document creation")

# Integrating Steps
def main(pdf_path, template_path, output_path, api_key_file, target_language="en"):
    api_key = read_api_key(api_key_file)
    gemini = gemini_api_auth(api_key)
    text, figures, tables = extract_text_from_pdf(pdf_path)
    corrected_text = correct_text(text, gemini)
    translated_text = translate_text(corrected_text, gemini, target_language)
    labeled_text = label_figures_text(translated_text, figures)
    create_word_document(template_path, labeled_text, figures, output_path)

# Run the main function
pdf_path = "input.pdf"
template_path = "Temp_1.docx"
output_path = "output.docx"
api_key_file = "apikey.txt"  # Path to your API key file
main(pdf_path, template_path, output_path, api_key_file, target_language="en")